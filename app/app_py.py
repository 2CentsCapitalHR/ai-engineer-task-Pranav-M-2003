# -*- coding: utf-8 -*-
"""app.py

Original file is located at
    https://colab.research.google.com/drive/1SBPer6gS1vFFdr9EoeSf0PklpSnEbhYv
"""

# !pip install gradio python-docx

import gradio as gr
import docx
import json
import os
import time

# --- Core Application Logic ---

def simulate_llm_analysis(doc_name):
    """
    This function simulates an AI/LLM analysis based on document name.
    In a real-world scenario, this would involve a call to an LLM API (like GPT,
    Gemini, etc.) with the document's text and a RAG context.
    """
    analysis = {
        "doc_type": "Unknown",
        "issues": []
    }

    lower_doc_name = doc_name.lower()

    if 'articles' in lower_doc_name or 'aoa' in lower_doc_name:
        analysis["doc_type"] = "Articles of Association"
        analysis["issues"].append({
            "section": "Jurisdiction Clause (Simulated)",
            "issue": "Jurisdiction clause does not explicitly specify 'ADGM Courts'.",
            "severity": "High",
            "suggestion": "Update the jurisdiction clause to 'ADGM Courts' to comply with ADGM regulations."
        })
    elif 'board' in lower_doc_name and 'resolution' in lower_doc_name:
         analysis["doc_type"] = "Board Resolution"
         analysis["issues"].append({
            "section": "Signatory Section (Simulated)",
            "issue": "Missing date on the signatory page.",
            "severity": "Medium",
            "suggestion": "Ensure all signatory fields, including the date, are properly filled out."
        })
    elif 'memo' in lower_doc_name or 'moa' in lower_doc_name:
        analysis["doc_type"] = "Memorandum of Association"
    elif 'ubo' in lower_doc_name:
        analysis["doc_type"] = "UBO Declaration Form"
    elif 'register' in lower_doc_name:
        analysis["doc_type"] = "Register of Members and Directors"
    else:
        analysis["doc_type"] = "General Document"
        analysis["issues"].append({
            "section": "General Compliance (Simulated)",
            "issue": "Document format does not appear to match standard ADGM templates.",
            "severity": "Low",
            "suggestion": "Cross-reference with the official ADGM templates to ensure full compliance and formatting."
        })
    return analysis

def analyze_documents(files, progress=gr.Progress()):
    """
    The main function that processes uploaded files, analyzes them,
    and generates the final report and commented documents.
    """
    if not files:
        raise gr.Error("Please upload at least one document.")

    # Define the checklist for a standard company incorporation
    incorporation_checklist = {
        "Articles of Association (AoA)": False,
        "Memorandum of Association (MoA/MoU)": False,
        "Board Resolution Templates": False,
        "Shareholder Resolution Templates": False,
        "Incorporation Application Form": False,
        "UBO Declaration Form": False,
        "Register of Members and Directors": False
    }

    all_issues = []
    reviewed_file_paths = []

    # Create a temporary directory for output files
    output_dir = "reviewed_docs"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    total_files = len(files)
    for i, file in enumerate(files):
        progress(i / total_files, desc=f"Analyzing {os.path.basename(file.name)}...")

        doc_name = os.path.basename(file.name)

        # --- Document Parsing and Checklist Update ---
        try:
            doc = docx.Document(file.name)
            # In a real app, you'd extract text: full_text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            # For this demo, we create a blank doc to add comments to
            doc = docx.Document()
            doc.add_paragraph(f"Original content of {doc_name} would be here.")
            doc.add_paragraph(f"(Could not fully parse original, but analysis is proceeding based on filename)")

        for key in incorporation_checklist:
            simplified_key = key.split('(')[0].strip().lower()
            if simplified_key in doc_name.lower():
                incorporation_checklist[key] = True

        # --- AI Analysis (Simulated) ---
        analysis_result = simulate_llm_analysis(doc_name)

        if analysis_result.get("issues"):
            all_issues.extend([{"document": doc_name, **issue} for issue in analysis_result["issues"]])

            # --- Insert Comments into DOCX ---
            doc.add_section()
            doc.add_heading('--- AI Analysis & Comments ---', level=2)
            for issue in analysis_result["issues"]:
                doc.add_paragraph(
                    f"COMMENT: In section '{issue['section']}', an issue was found: '{issue['issue']}' (Severity: {issue['severity']}).",
                    style='Intense Quote'
                )
                doc.add_paragraph(f"SUGGESTION: {issue['suggestion']}")
                doc.add_paragraph()

        # --- Save Reviewed Document ---
        reviewed_file_path = os.path.join(output_dir, f"reviewed_{doc_name}")
        doc.save(reviewed_file_path)
        reviewed_file_paths.append(reviewed_file_path)
        time.sleep(0.5) # Simulate processing time

    # --- Generate Final Report ---
    progress(1, desc="Generating final report...")

    # Checklist Verification
    missing_docs = [doc for doc, found in incorporation_checklist.items() if not found]
    checklist_summary = f"It appears you're trying to incorporate a company in ADGM. You have uploaded {len(files)} out of {len(incorporation_checklist)} required documents."
    if missing_docs:
        checklist_summary += f"\n\nThe missing document appears to be: '{missing_docs[0]}'."
    else:
        checklist_summary += "\n\nAll required documents for incorporation appear to be present."

    # Structured JSON Report
    report = {
        "process": "Company Incorporation",
        "documents_uploaded": len(files),
        "required_documents": len(incorporation_checklist),
        "missing_document": missing_docs[0] if missing_docs else "None",
        "issues_found": all_issues
    }

    return checklist_summary, report, reviewed_file_paths


# --- Gradio UI Definition ---
with gr.Blocks(theme=gr.themes.Soft(), title="ADGM Corporate Agent") as demo:
    gr.Markdown(
        """
        # ADGM-Compliant Corporate Agent
        Upload your `.docx` legal documents for review, validation, and compliance checking within the ADGM jurisdiction.
        """
    )

    with gr.Row():
        with gr.Column(scale=1):
            file_input = gr.File(
                label="Upload .docx Documents",
                file_count="multiple",
                file_types=[".docx"]
            )
            analyze_button = gr.Button("Analyze Documents", variant="primary")
            gr.Markdown(
                """
                **Note:** This is a demonstration. The AI analysis is simulated based on filenames.
                The agent will identify documents, check against an incorporation checklist, flag simulated issues,
                and generate downloadable `.docx` files with comments appended.
                """
            )

        with gr.Column(scale=2):
            gr.Markdown("## Analysis Results")
            checklist_output = gr.Textbox(label="Checklist Verification Summary", lines=4)
            json_output = gr.JSON(label="Structured Analysis Report")
            reviewed_files_output = gr.File(label="Download Reviewed Documents")

    analyze_button.click(
        fn=analyze_documents,
        inputs=file_input,
        outputs=[checklist_output, json_output, reviewed_files_output]
    )

if __name__ == "__main__":
    print("Launching Gradio Interface...")
    print("Please create a folder named 'reviewed_docs' in the same directory if it doesn't exist.")
    demo.launch()

